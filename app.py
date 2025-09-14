import io, re
import pandas as pd
import streamlit as st

from pypdf import PdfReader
from docx import Document as DocxDocument
from rapidfuzz import fuzz, process

# ---------- OCR fallback (auto-ignored if binaries missing) ----------
try:
    from pdf2image import convert_from_bytes
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False

# ---------- Page ----------
st.set_page_config(page_title="AI KPI System", layout="wide")
st.title("AI KPI Extraction & Recommendations (Per BRD)")

# ---------- Theme (red + white) ----------
st.markdown(
    """
    <style>
    :root { --brand:#b91c1c; }
    .stTextInput>div>div>input, .stSelectbox>div>div>select, textarea.stTextArea {
        border:1.5px solid var(--brand) !important; border-radius:6px !important; background:#fff !important;
    }
    .stTextInput>div>div>input:focus, .stSelectbox>div>div>select:focus, textarea.stTextArea:focus {
        border:2px solid var(--brand) !important; outline:none !important; box-shadow:0 0 5px var(--brand) !important;
    }

    .kpi-header { background:#b91c1c; color:#fff; padding:10px 12px; border-radius:8px 8px 0 0; font-weight:700; }
    .table-head { background:#f8fafc; border:1px solid #f1f5f9; border-bottom:0; border-radius:8px 8px 0 0; padding:8px 12px; font-weight:700; }
    .table-body { border:1px solid #f1f5f9; border-top:0; border-radius:0 0 8px 8px; }
    .cell { padding:10px 12px; border-top:1px solid #f1f5f9; }

    .badge { display:inline-block; padding:3px 8px; border-radius:999px; font-size:12px; color:#fff; }
    .badge-pending { background:#9ca3af; }
    .badge-validated { background:#16a34a; }
    .badge-rejected { background:#b91c1c; }

    .pill { display:inline-block; margin-top:6px; padding:6px 10px; border-radius:8px; font-weight:600; }
    .pill-green { background:#16a34a; color:#fff; }
    .pill-red { background:#b91c1c; color:#fff; }
    .pill-gray { background:#e5e7eb; color:#111827; }

    button[data-testid="baseButton-secondary"]{
        background:#f3f4f6 !important; color:#111827 !important; border:1px solid #e5e7eb !important;
        border-radius:8px !important; padding:6px 12px !important;
    }
    button[data-testid="baseButton-secondary"]:hover { background:#e5e7eb !important; }
    </style>
    """,
    unsafe_allow_html=True
)

# ---------- Session ----------
if "projects" not in st.session_state:
    st.session_state["projects"] = {}     # {brd: {...}}
if "final_kpis" not in st.session_state:
    st.session_state["final_kpis"] = {}   # brd -> DataFrame

def _status_badge(s):
    cls = "badge-pending"
    if s == "Validated": cls = "badge-validated"
    elif s == "Rejected": cls = "badge-rejected"
    return f"<span class='badge {cls}'>{s}</span>"

def _action_pill(status):
    if status == "Validated": return "<span class='pill pill-green'>Validated</span>"
    if status == "Rejected":  return "<span class='pill pill-red'>Rejected</span>"
    return "<span class='pill pill-gray'>Pending</span>"

def _upsert_final(brd, row):
    df = st.session_state["final_kpis"].get(
        brd,
        pd.DataFrame(columns=["BRD","KPI Name","Source","Description","Owner/ SME","Target Value"])
    )
    df = pd.concat([df, pd.DataFrame([row])], ignore_index=True)
    df.drop_duplicates(subset=["KPI Name"], keep="last", inplace=True)
    st.session_state["final_kpis"][brd] = df

def _remove_from_final(brd, kpi_name):
    df = st.session_state["final_kpis"].get(
        brd,
        pd.DataFrame(columns=["BRD","KPI Name","Source","Description","Owner/ SME","Target Value"])
    )
    if not df.empty:
        df = df[df["KPI Name"] != kpi_name].reset_index(drop=True)
    st.session_state["final_kpis"][brd] = df

# ---------- File reading (PDF + OCR, DOCX) ----------
def read_text_from_bytes(data, name):
    lname = name.lower()
    bio = io.BytesIO(data)

    if lname.endswith(".pdf"):
        # native text
        try:
            reader = PdfReader(bio)
            txt = "\n".join((p.extract_text() or "") for p in reader.pages)
            if txt and txt.strip():
                return txt
        except Exception:
            pass
        # OCR fallback
        if OCR_AVAILABLE:
            try:
                imgs = convert_from_bytes(data)
                parts = []
                for im in imgs:
                    t = pytesseract.image_to_string(im, lang="eng")
                    if t and t.strip():
                        parts.append(t)
                return "\n".join(parts)
            except Exception:
                return ""
        return ""

    if lname.endswith(".docx"):
        try:
            doc = DocxDocument(bio)
        except Exception:
            return ""
        parts = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t: parts.append(t)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t: parts.append(t)
        return "\n".join(parts)

    try:
        return data.decode(errors="ignore")
    except Exception:
        return ""

def read_uploaded(file):
    return read_text_from_bytes(file.read(), file.name)

# ---------- Domain & Topic inference ----------
DOMAIN_HINTS = {
    "hr": [
        "employee","attrition","turnover","recruitment","hiring","retention",
        "satisfaction","absenteeism","time to fill","job description","resume",
        "cv","parsing","ats","matching","screening"
    ],
    "sales": ["pipeline","deal","quota","win rate","opportunity","lead"],
    "marketing": ["campaign","cpl","cac","ctr","impressions","engagement"],
    "finance": ["revenue","margin","cash","roi","ebitda"],
}
def infer_domain(text):
    low = text.lower()
    scores = {d: sum(low.count(k) for k in ks) for d, ks in DOMAIN_HINTS.items()}
    d = max(scores, key=scores.get)
    return d if scores[d] > 0 else "hr"

HR_TOPIC_KEYWORDS = {
    "attrition_retention": ["attrition","turnover","retention","churn","stay interview","exit interview"],
    "jd_ats": ["job description","jd","resume","cv","parsing","ats","matching","skills extraction","screening","latency","concurrent users"],
    "workforce_planning": ["headcount","workforce planning","manpower","capacity","utilization","vacancy","forecast"],
    "recruiting": ["recruiting","sourcing","hiring","offer","candidate","interview","requisition"],
    "learning_development": ["training","learning","l&d","course","skill","upskilling","certification","completion"],
    "engagement_culture": ["engagement","enps","nps","pulse","survey","satisfaction","morale","culture","recognition"],
}
def infer_hr_topic(text):
    low = text.lower()
    best, score = None, 0
    for topic, kws in HR_TOPIC_KEYWORDS.items():
        s = sum(low.count(k) for k in kws)
        if s > score:
            best, score = topic, s
    return best or "recruiting"

# ---------- Text preprocessing ----------
SPLIT_TOKENS = r"[•\u2022\-\–\—\·]|(?:^\s*\d+[\.\)])"
def preprocess_text(raw):
    lines = [l.strip() for l in raw.splitlines()]
    fused, buf = [], ""
    for l in lines:
        if not l:
            if buf: fused.append(buf.strip()); buf = ""
            continue
        if buf and not re.search(r"[\.!\?;:]$", buf) and not re.match(r"^\s*-\s*", l):
            buf = f"{buf} {l}"
        else:
            if buf: fused.append(buf.strip())
            buf = l
    if buf: fused.append(buf.strip())

    bullets = []
    for chunk in fused:
        for s in re.split(SPLIT_TOKENS, chunk):
            s = re.sub(r"\s+", " ", s).strip(" -•:\t")
            if s: bullets.append(s)

    sents = []
    for b in bullets:
        for sent in re.split(r"(?<=[\.\?!])\s+(?=[A-Z])", b):
            s = sent.strip()
            if s: sents.append(s)
    return sents

# ---------- KPI patterns & extraction ----------
TARGET_PERCENT = r"\b(?:<|>|≤|≥)?\s*\d{1,3}(?:\.\d+)?\s*%\b"
TARGET_RATIO   = r"\b\d+(?:\.\d+)?\s*/\s*\d+\b"
TARGET_TIME    = r"\b(?:in|within|by)\s+\d+\s*(?:days?|weeks?|months?|quarters?|years?)\b"
TARGET_SECS    = r"\b\d+(?:\.\d+)?\s*(?:ms|milliseconds|s|sec|secs|seconds)\b"
TARGET_UNDER_SECS = r"\b(?:under|within|less than|<|≤)\s*\d+(?:\.\d+)?\s*(?:ms|milliseconds|s|sec|secs|seconds)\b"
TARGET_USERS   = r"\b\d{1,5}(?:\.\d+)?\s*(?:concurrent\s+users|users)\b"
TARGET_GENERIC = r"(?:<|>|≤|≥)\s*\d+(?:\.\d+)?"

def find_targets(s):
    hits = []
    for pat in (TARGET_PERCENT, TARGET_RATIO, TARGET_TIME, TARGET_UNDER_SECS, TARGET_SECS, TARGET_USERS, TARGET_GENERIC):
        for m in re.finditer(pat, s, flags=re.I):
            hits.append(m.group(0).strip())
    return " | ".join(dict.fromkeys(hits))

KPI_CANON = {
    # JD/ATS & platform KPIs
    "JD Generation Time": [
        r"\bgenerate\b.{0,40}\bjd\b.{0,40}\b(sec|seconds|ms|milliseconds|time|latency)\b",
        r"\bjd (generation|creation) time\b"
    ],
    "Bias Flag Rate": [r"\bbias (flag|detection)\b", r"\bnon[- ]inclusive language\b"],
    "JD Tool Adoption Rate": [r"\badoption rate\b", r"\busage rate\b", r"\b% of (hiring managers|users) using\b"],
    "JD Approval Rate": [r"\bapproval rate\b", r"\bapproved without major edits\b"],
    "Repository Compliance": [r"\bversion control\b", r"\bapproval logs\b", r"\brepository\b"],
    "System Uptime": [r"\buptime\b", r"\bavailability\b", r"\b99\.\d{1,2}% availability\b"],
    "Concurrent Users Supported": [r"\bconcurrent users\b"],
    "JD Parsing Accuracy": [r"\bjd parsing\b", r"\bjob description parsing\b"],
    "Skills Extraction Precision": [r"\bskills extraction\b", r"\bskill extraction\b"],
    "Resume Matching Accuracy": [r"\bresume matching\b", r"\bcv matching\b"],
    "JD-Resume Match Score": [r"\bmatch score\b"],
    "Average Screening Time": [r"\bscreening time\b", r"\bscreen time\b"],
    "Response Latency": [r"\blatency\b", r"\bresponse time\b"],

    # Core HR
    "Voluntary Attrition Rate": [r"\bvoluntary attrition\b", r"\bvoluntary turnover\b"],
    "Involuntary Attrition Rate": [r"\binvoluntary attrition\b", r"\binvoluntary turnover\b"],
    "Employee Retention Rate": [r"\bretention rate\b", r"\bemployee retention\b"],
    "First Year Attrition Rate": [r"\bfirst[- ]year attrition\b"],
    "Average Tenure": [r"\baverage tenure\b", r"\bavg tenure\b"],
    "Internal Mobility Rate": [r"\binternal mobility\b"],
    "Absenteeism Rate": [r"\babsenteeism\b"],
    "Employee Satisfaction Score": [r"\bemployee satisfaction\b", r"\bsatisfaction score\b"],
    "Employee NPS": [r"\bemployee nps\b", r"\benps\b", r"\bnet promoter score\b"],
    "Engagement Score": [r"\bengagement score\b"],
    "Time to Fill": [r"\btime to fill\b"],
    "Time to Hire": [r"\btime to hire\b"],
    "Offer Acceptance Rate": [r"\boffer acceptance\b"],
    "Quality of Hire": [r"\bquality of hire\b"],
    "Cost per Hire": [r"\bcost per hire\b"],
}

EXCLUDE_PHRASES = [
    "business requirements document","brd","document","project","model","introduction","purpose","scope",
    "assumptions","out of scope","table of contents","revision history","version","author","date","appendix"
]
METRIC_WORDS = re.compile(
    r"\b(rate|ratio|score|index|time|latency|throughput|tenure|utilization|accuracy|precision|recall|f1|cost|uptime|availability|concurrent|match)\b",
    re.I,
)

def canonical_from_text(s):
    low = s.lower()
    for canon, pats in KPI_CANON.items():
        for pat in pats:
            if re.search(pat, low):
                return canon
    return None

def is_probably_kpi(line):
    low = line.lower()
    if any(ph in low for ph in EXCLUDE_PHRASES): return False
    if canonical_from_text(low): return True
    if re.search(r"\b(shall|should|must|system)\b", low) and METRIC_WORDS.search(low): return True
    if METRIC_WORDS.search(low) and (find_targets(low) or re.search(r"\b(kpi|metric)\b", low)): return True
    return False

KPI_SEEDS = list(KPI_CANON.keys())

def normalize_name(raw_line):
    canon = canonical_from_text(raw_line)
    if canon: return canon
    guess = re.split(r"[:\-–]| \(", raw_line, maxsplit=1)[0].strip()
    if len(guess.split()) > 6 and not METRIC_WORDS.search(guess):
        guess = " ".join(guess.split()[:6])
    match = process.extractOne(guess, KPI_SEEDS, scorer=fuzz.WRatio)
    if match and match[1] >= 85: return match[0]
    return guess.title()

def dedup_rows(df):
    if df.empty: return df
    def _norm(k): return re.sub(r"[^a-z0-9]+", " ", str(k).lower()).strip()
    best = {}
    for _, r in df.iterrows():
        key = _norm(r["KPI Name"])
        cur = best.get(key)
        if cur is None or (not cur["Target Value"] and r["Target Value"]):
            best[key] = r
    return pd.DataFrame(best.values()).reset_index(drop=True)

def extract_kpis(text):
    cols = ["KPI Name", "Description", "Target Value", "Status"]
    rows = []
    for ln in preprocess_text(text):
        if not is_probably_kpi(ln):
            continue
        name = normalize_name(ln)
        targets = find_targets(ln)
        desc = ln if len(ln) <= 240 else ln[:237] + "…"
        rows.append({"KPI Name": name, "Description": desc, "Target Value": targets, "Status": "Pending"})
    df = pd.DataFrame(rows, columns=cols)
    df = dedup_rows(df)
    return df if not df.empty else pd.DataFrame(columns=cols)

# ---------- Description generator for Recommended KPIs ----------
RECOMMENDED_DESC = {
    # JD/ATS
    "JD Parsing Accuracy": "Percentage of job descriptions correctly parsed into structured fields.",
    "Skills Extraction Precision": "Precision of the skills extraction model on annotated JD/resume samples.",
    "Resume Matching Accuracy": "Share of candidate-resume matches correctly ranked in the top results.",
    "JD-Resume Match Score": "Average relevance score between a job description and shortlisted resumes.",
    "Average Screening Time": "Average time recruiters take to screen a candidate profile.",
    "Response Latency": "Average system response time for JD/ATS actions and searches.",
    "Concurrent Users Supported": "Maximum number of users supported concurrently without degradation.",
    "System Uptime": "Percent of time the platform is available to users in a period.",
    # Core HR
    "Voluntary Attrition Rate": "Share of employees who leave voluntarily within the period.",
    "Involuntary Attrition Rate": "Share of separations initiated by the employer within the period.",
    "Employee Retention Rate": "Percent of employees retained over a 12-month period.",
    "First Year Attrition Rate": "Share of new hires leaving within their first 12 months.",
    "Average Tenure": "Average length of service for active employees.",
    "Internal Mobility Rate": "Percent of roles filled by internal candidates.",
    "Absenteeism Rate": "Unplanned absence days as a percentage of total scheduled days.",
    "Employee Satisfaction Score": "Average score from periodic satisfaction/engagement surveys.",
    "Employee NPS": "Employee Net Promoter Score from pulse or annual surveys.",
    "Time to Fill": "Average number of days between job posting and offer acceptance.",
    "Time to Hire": "Average number of days between candidate application and offer acceptance.",
    "Offer Acceptance Rate": "Percent of offers accepted by candidates.",
    "Quality of Hire": "Composite index of performance, retention, and cultural fit of new hires.",
    "Cost per Hire": "Average end-to-end recruiting cost to make a hire.",
}

def generate_description(kpi_name: str, topic: str | None) -> str:
    # Priority: hard-coded mapping; otherwise short generic template
    if kpi_name in RECOMMENDED_DESC:
        return RECOMMENDED_DESC[kpi_name]
    # Fallback short generic
    base = kpi_name.rstrip(".")
    return f"Standard definition for {base} measured consistently across the period."

# ---------- Recommendations ----------
TOPIC_KPIS = {
    "attrition_retention": [
        "Voluntary Attrition Rate","Involuntary Attrition Rate","Employee Retention Rate",
        "First Year Attrition Rate","Average Tenure","Internal Mobility Rate",
        "Stay Interview Coverage","Exit Interview Completion Rate","Regretted Attrition Rate"
    ],
    "jd_ats": [
        "JD Parsing Accuracy","Skills Extraction Precision","Resume Matching Accuracy",
        "JD-Resume Match Score","Automation Rate","Throughput Per Hour",
        "Average Screening Time","Recruiter Productivity","False Positive Match Rate","SLA Compliance"
    ],
    "workforce_planning": [
        "Headcount Growth","Vacancy Rate","Time to Backfill","Capacity Utilization",
        "Forecast Accuracy","Bench Strength Index","Span of Control"
    ],
    "recruiting": [
        "Time to Fill","Time to Hire","Offer Acceptance Rate","Quality of Hire",
        "Cost per Hire","Candidate Drop-off Rate","Interview to Offer Ratio"
    ],
    "learning_development": [
        "Training Completion Rate","Training Effectiveness Score","Certification Pass Rate",
        "Learning Hours per Employee","Time to Competency","Skills Coverage Index"
    ],
    "engagement_culture": [
        "Engagement Score","Employee NPS","Participation Rate (Surveys)",
        "Recognition Frequency","Manager Feedback Response Time","eNPS Promoter Ratio"
    ],
}
FALLBACK_KPIS = {
    "hr": [
        "Offer Acceptance Rate","Absenteeism Rate","Training Completion Rate","Employee NPS",
        "Internal Mobility Rate","Quality of Hire","Cost per Hire","Diversity Ratio",
        "Vacancy Rate","First Year Attrition Rate"
    ],
    "sales": ["Win Rate","Lead Conversion","Quota Attainment","Average Deal Size","Sales Cycle Length"],
    "marketing": ["CTR","CAC","CPL","Brand Awareness Index","Email Open Rate"],
    "finance": ["Operating Margin","EBITDA Margin","Cash Burn","Runway Months","DSO"],
}

def recommend(domain, existing, topic=None, raw_text=""):
    existing_l = {e.lower() for e in existing}
    if domain == "hr":
        topic = topic or infer_hr_topic(raw_text)
        pool = TOPIC_KPIS.get(topic, FALLBACK_KPIS["hr"])
    else:
        pool = FALLBACK_KPIS.get(domain, [])
    out, seen = [], set()
    for k in pool:
        if k.lower() in existing_l or k.lower() in seen: continue
        out.append(k); seen.add(k.lower())
        if len(out) >= 12: break
    return out, topic

# ---------- UI helpers ----------
def _table_head(cols, headers):
    st.markdown(
        "<div class='table-head' style='display:grid;grid-template-columns:" +
        " ".join(cols) + ";'>" +
        "".join([f"<div>{h}</div>" for h in headers]) +
        "</div>", unsafe_allow_html=True
    )
    st.markdown("<div class='table-body'>", unsafe_allow_html=True)

def _table_tail():
    st.markdown("</div>", unsafe_allow_html=True)

def render_extracted_table(brd, df, key_prefix):
    if df.empty:
        st.caption("No extracted KPIs.")
        return df
    _table_head(["2fr","3fr","1fr","0.9fr","1.6fr"], ["KPI Name","Description","Target Value","Status","Actions"])
    updated = []
    for i, r in df.iterrows():
        c1, c2, c3, c4, c5 = st.columns([2,3,1,0.9,1.6])
        with c1: st.markdown(f"<div class='cell'><b>{r['KPI Name']}</b></div>", unsafe_allow_html=True)
        with c2: st.markdown(f"<div class='cell'>{r['Description']}</div>", unsafe_allow_html=True)
        with c3: target_val = st.text_input("", value=r["Target Value"], key=f"{key_prefix}_target_{i}")
        with c4: st.markdown(f"<div class='cell'>{_status_badge(r['Status'])}</div>", unsafe_allow_html=True)
        with c5:
            colB, colC = st.columns([1,1])
            if colB.button("Validate", key=f"{key_prefix}_ok_{i}"):
                r["Status"] = "Validated"
                _upsert_final(brd, {
                    "BRD": brd,
                    "KPI Name": r["KPI Name"],
                    "Source": "Extracted",
                    "Description": r["Description"],
                    "Owner/ SME": "",
                    "Target Value": target_val
                })
            if colC.button("Reject", key=f"{key_prefix}_rej_{i}"):
                r["Status"] = "Rejected"
                _remove_from_final(brd, r["KPI Name"])
            st.markdown(_action_pill(r["Status"]), unsafe_allow_html=True)
        updated.append({"KPI Name": r["KPI Name"], "Description": r["Description"], "Target Value": target_val, "Status": r["Status"]})
    _table_tail()
    return pd.DataFrame(updated, columns=list(df.columns))

def render_recommended_table(brd, df, key_prefix):
    if df.empty:
        st.caption("No recommendations.")
        return df
    # KPI Name is READ-ONLY now; Description is editable
    _table_head(
        ["2fr","2.5fr","1fr","1fr","0.8fr","1.6fr"],
        ["KPI Name","Description","Owner/ SME","Target Value","Status","Actions"]
    )
    updated = []
    for i, r in df.iterrows():
        c1, c2, c3, c4, c5, c6 = st.columns([2,2.5,1,1,0.8,1.6])
        with c1: st.markdown(f"<div class='cell'><b>{r['KPI Name']}</b></div>", unsafe_allow_html=True)
        with c2: desc_val = st.text_input("", value=r.get("Description",""), key=f"{key_prefix}_desc_{i}")
        with c3: owner_val = st.text_input("", value=r.get("Owner/ SME",""), key=f"{key_prefix}_owner_{i}")
        with c4: target_val = st.text_input("", value=r.get("Target Value",""), key=f"{key_prefix}_target_{i}")
        with c5: st.markdown(f"<div class='cell'>{_status_badge(r['Status'])}</div>", unsafe_allow_html=True)
        with c6:
            colB, colC = st.columns([1,1])
            if colB.button("Validate", key=f"{key_prefix}_ok_{i}"):
                r["Status"] = "Validated"
                _upsert_final(brd, {
                    "BRD": brd,
                    "KPI Name": r["KPI Name"],
                    "Source": "Recommended",
                    "Description": desc_val,
                    "Owner/ SME": owner_val,
                    "Target Value": target_val
                })
            if colC.button("Reject", key=f"{key_prefix}_rej_{i}"):
                r["Status"] = "Rejected"
                _remove_from_final(brd, r["KPI Name"])
            st.markdown(_action_pill(r["Status"]), unsafe_allow_html=True)

        updated.append({
            "KPI Name": r["KPI Name"],
            "Description": desc_val,
            "Owner/ SME": owner_val,
            "Target Value": target_val,
            "Status": r["Status"]
        })
    _table_tail()
    return pd.DataFrame(updated, columns=list(df.columns))

# ---------- Manual KPI adder (Recommended section) ----------
def manual_kpi_adder(brd, topic):
    st.markdown("#### Add KPI manually")
    with st.form(key=f"manual_add_{brd}", clear_on_submit=True):
        c1, c2 = st.columns([2,2])
        kpi_name = c1.text_input("KPI Name *", value="")
        # pre-fill description from generator as you type (best effort)
        suggested_desc = generate_description(kpi_name.strip(), topic) if kpi_name.strip() else ""
        desc     = c2.text_input("Description", value=suggested_desc)
        c3, c4 = st.columns([1,1])
        owner    = c3.text_input("Owner/ SME", value="")
        target   = c4.text_input("Target Value", value="")
        add = st.form_submit_button("Add KPI")
    if add:
        if not kpi_name.strip():
            st.warning("Please enter a KPI Name.")
            return
        rec_df = st.session_state["projects"][brd]["recommended"]
        all_existing = set([x.lower() for x in rec_df["KPI Name"].astype(str).tolist()])
        ext_df = st.session_state["projects"][brd]["extracted"]
        all_existing |= set([x.lower() for x in ext_df["KPI Name"].astype(str).tolist()])
        if kpi_name.strip().lower() in all_existing:
            st.warning("That KPI already exists in this BRD.")
            return
        new_row = {
            "KPI Name": kpi_name.strip(),
            "Description": (desc.strip() or generate_description(kpi_name.strip(), topic)),
            "Owner/ SME": owner.strip(),
            "Target Value": target.strip(),
            "Status": "Pending"
        }
        rec_df = pd.concat([rec_df, pd.DataFrame([new_row])], ignore_index=True)
        st.session_state["projects"][brd]["recommended"] = rec_df
        st.success("KPI added to Recommended.")
        st.rerun()

# ---------- Pipeline per file ----------
def process_file(file):
    text = read_uploaded(file)
    domain = infer_domain(text)
    topic  = infer_hr_topic(text) if domain == "hr" else None

    extracted = extract_kpis(text)
    existing = extracted["KPI Name"].astype(str).tolist() if not extracted.empty else []
    recs, topic = recommend(domain, existing, topic=topic, raw_text=text)

    # Build recommended with auto-descriptions
    rows = []
    for r in recs:
        rows.append({
            "KPI Name": r,
            "Description": generate_description(r, topic),
            "Owner/ SME": "",
            "Target Value": "",
            "Status": "Pending"
        })
    recommended = pd.DataFrame(rows, columns=["KPI Name", "Description", "Owner/ SME", "Target Value", "Status"])

    st.session_state["projects"][file.name] = {
        "domain": domain, "topic": topic,
        "extracted": extracted, "recommended": recommended
    }
    st.session_state["final_kpis"].setdefault(
        file.name,
        pd.DataFrame(columns=["BRD","KPI Name","Source","Description","Owner/ SME","Target Value"])
    )

# ---------- Main ----------
uploads = st.file_uploader("Upload BRDs", type=["pdf","docx","txt"], accept_multiple_files=True)

if st.button("Process BRDs"):
    if not uploads:
        st.warning("Please upload at least one file")
    else:
        for f in uploads:
            process_file(f)
        count = len(uploads)
        st.success(f"✅ Processed {count} BRD{'s' if count != 1 else ''} successfully")

# ---------- Render per BRD ----------
for fname, proj in st.session_state["projects"].items():
    topic_text = f" — Topic: **{proj.get('topic','').replace('_',' ').title()}**" if proj.get("topic") else ""
    st.markdown(f"## 📄 {fname} — Domain: **{proj['domain'].upper()}**{topic_text}")

    st.subheader("Extracted KPIs")
    proj["extracted"] = render_extracted_table(fname, proj["extracted"], key_prefix=f"ext_{fname}")

    st.subheader("Recommended KPIs")
    # Manual add FIRST so new rows appear instantly
    manual_kpi_adder(fname, proj.get("topic"))
    proj["recommended"] = render_recommended_table(fname, proj["recommended"], key_prefix=f"rec_{fname}")

    st.markdown("<div class='kpi-header'>Finalized KPIs (This BRD)</div>", unsafe_allow_html=True)
    final_df = st.session_state["final_kpis"].get(fname, pd.DataFrame())
    if final_df.empty:
        st.caption("No validated KPIs yet for this BRD. Use Validate above.")
    else:
        show = final_df[["KPI Name","Source","Owner/ SME","Target Value","Description"]].sort_values("KPI Name")
        st.dataframe(show, use_container_width=True, hide_index=True)
